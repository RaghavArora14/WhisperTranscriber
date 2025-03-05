import streamlit as st
import os
import whisper
import tempfile
from datetime import datetime
from docx import Document
from audio_recorder_streamlit import audio_recorder
import tkinter as tk
from tkinter import filedialog

# Initialize Whisper model with language support
@st.cache_resource
def load_model():
    return whisper.load_model("base")

model = load_model()

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.mp3', '.wav', '.mp4', '.mkv', '.mov', '.flv', '.aac', '.m4a'}

# Supported languages
LANGUAGES = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "ja": "Japanese",
    "zh": "Chinese",
    "hi": "Hindi",
    "ru": "Russian",
    "pt": "Portuguese",
    "ar": "Arabic",
    "ko": "Korean",
    # Add more languages as needed
}

# Function to create transcription document
def create_transcription_doc(transcriptions, output_path):
    doc = Document()
    for file_name, created_date, text in transcriptions:
        doc.add_heading(f"{file_name} - {created_date.strftime('%Y-%m-%d %H:%M:%S')}", level=1)
        doc.add_paragraph(text)
        doc.add_page_break()
    doc.save(output_path)

# Function to select directory
def select_directory():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    root.wm_attributes('-topmost', 1)  # Keep the dialog on top
    folder = filedialog.askdirectory()
    return folder

# Main app function
def main():
    st.title("🎙️ Audio Transcription System")
    
    # Sidebar for settings
    with st.sidebar:
        st.header("Settings")
        model_size = st.selectbox(
            "Model Size",
            ("base", "small", "medium", "large"),
            index=0
        )
        
        # Language selection
        selected_lang = st.selectbox(
            "Select Language",
            options=list(LANGUAGES.values()),
            index=0
        )
        # Get language code
        lang_code = [k for k, v in LANGUAGES.items() if v == selected_lang][0]

    # Tab layout
    tab1, tab2 = st.tabs(["File Transcription", "Live Recording"])
    
    with tab1:
        st.header("File Transcription")
        
        if st.button("Select Directory"):
            selected_dir = select_directory()
            if selected_dir:
                st.session_state['selected_dir'] = selected_dir
        
        if 'selected_dir' in st.session_state:
            st.write(f"Selected Directory: {st.session_state['selected_dir']}")
            
            if st.button("Transcribe Files"):
                directory = st.session_state['selected_dir']
                if not os.path.isdir(directory):
                    st.error("Invalid directory path")
                else:
                    with st.spinner("Processing files..."):
                        transcriptions = []
                        all_text = ""
                        for root, _, files in os.walk(directory):
                            for file in files:
                                file_path = os.path.join(root, file)
                                ext = os.path.splitext(file_path)[1].lower()
                                if ext in SUPPORTED_EXTENSIONS:
                                    try:
                                        created_date = datetime.fromtimestamp(os.path.getctime(file_path))
                                        result = model.transcribe(file_path, language=lang_code)
                                        file_text = f"{file} - {created_date.strftime('%Y-%m-%d %H:%M:%S')}\n\n{result['text']}\n\n"
                                        all_text += file_text
                                        transcriptions.append((file, created_date, result['text']))
                                    except Exception as e:
                                        st.error(f"Error processing {file}: {str(e)}")
                        
                        if transcriptions:
                            output_path = os.path.join(directory, "transcriptions.docx")
                            create_transcription_doc(transcriptions, output_path)
                            st.success(f"Transcription completed! Processed {len(transcriptions)} files")
                            
                            # Show transcription text
                            st.subheader("Transcription Results")
                            st.text_area("Transcription Text", all_text, height=400)
                            
                            # Download options
                            col1, col2 = st.columns(2)
                            with col1:
                                with open(output_path, "rb") as file:
                                    st.download_button(
                                        label="Download as DOCX",
                                        data=file,
                                        file_name="transcriptions.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                            with col2:
                                st.download_button(
                                    label="Download as TXT",
                                    data=all_text,
                                    file_name="transcriptions.txt",
                                    mime="text/plain"
                                )
                        else:
                            st.info("No supported files found for transcription")

    with tab2:
        st.header("Live Recording")
        st.write("Record audio directly from your microphone")
        
        # Audio recorder
        audio_bytes = audio_recorder()
        
        if audio_bytes:
            st.audio(audio_bytes, format="audio/wav")
            
            if st.button("Transcribe Recording"):
                with st.spinner("Transcribing..."):
                    try:
                        # Save temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                            tmp_file.write(audio_bytes)
                            tmp_path = tmp_file.name
                        
                        # Transcribe
                        result = model.transcribe(tmp_path, language=lang_code)
                        os.remove(tmp_path)
                        
                        st.success("Transcription completed!")
                        transcription_text = result['text']
                        
                        # Show transcription
                        st.text_area("Transcription", transcription_text, height=300)
                        
                        # Download options
                        col1, col2 = st.columns(2)
                        with col1:
                            # Create temporary DOCX file
                            doc = Document()
                            doc.add_paragraph(transcription_text)
                            doc_path = os.path.join(tempfile.gettempdir(), "transcription.docx")
                            doc.save(doc_path)
                            with open(doc_path, "rb") as file:
                                st.download_button(
                                    label="Download as DOCX",
                                    data=file,
                                    file_name="transcription.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        with col2:
                            st.download_button(
                                label="Download as TXT",
                                data=transcription_text,
                                file_name="transcription.txt",
                                mime="text/plain"
                            )
                    except Exception as e:
                        st.error(f"Error during transcription: {str(e)}")

if __name__ == "__main__":
    main()